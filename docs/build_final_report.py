#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
3단계 최종 보고서(.docx) 생성기.
- 2단계 테스트 보고서 docx를 복제해 표지·머리글·푸터·로고·스타일을 그대로 계승.
- 표지 정보표를 최종 보고서용으로 수정.
- 본문(섹션 1부터)을 전부 제거하고 docs/final-report.md 내용을 템플릿 스타일로 재구성.
- mermaid 코드블록은 [그림] 캡션 + 코드 박스로 표기, 스크린샷 이미지는 삽입.
- 릴리즈 버전은 client-v0.1.2 기준.
"""
import re, os
from docx import Document
from docx.shared import Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.table import Table

SRC = "/opt/se_project/SE-final-project/docs/submissions/phase2/소프트웨어공학_2단계_테스트보고서.docx"
MD  = "/opt/se_project/SE-final-project/docs/final-report.md"
OUT = "/opt/se_project/SE-final-project/docs/submissions/phase3/소프트웨어공학_3단계_최종보고서.docx"
SHOT_DIR = "/opt/se_project/SE-final-project/client/screenshots"

# 템플릿 색상
C_NUM   = RGBColor(0x00, 0x3F, 0xD2)   # 섹션 번호
C_TITLE = RGBColor(0x24, 0x4D, 0xAE)   # 섹션 제목 / 표 헤더 배경
C_TEXT  = RGBColor(0x00, 0x00, 0x00)
HDR_FILL = "244DAE"
ROW_ALT  = "F2F5FB"
FONT_KO  = "Pretendard Variable"
FONT_MONO= "D2Coding"

doc = Document(SRC)

# ── 1. 표지 텍스트 수정 ──────────────────────────────
def set_para_text_keep_format(p, new_text):
    """단락의 첫 run 텍스트만 교체, 나머지 run 제거 (서식 유지)."""
    if not p.runs:
        p.add_run(new_text); return
    p.runs[0].text = new_text
    for r in p.runs[1:]:
        r._r.getparent().remove(r._r)

for p in doc.paragraphs[:6]:
    t = p.text.strip()
    if t == "테스트 보고서":
        set_para_text_keep_format(p, "최종 보고서")
    elif t == "Test Report":
        set_para_text_keep_format(p, "Final Report")
    elif t == "소프트웨어 공학 조별과제 2단계 (Phase 2) 제출용":
        set_para_text_keep_format(p, "소프트웨어 공학 조별과제 3단계 (Phase 3) 제출용")

# 표지 정보표(첫 표) 값 수정
cover = doc.tables[0]
cover_map = {
    "제출 단계": "3단계 최종 보고서",
    "측정 기준일": "2026-06-18",
}
for row in cover.rows:
    key = row.cells[0].text.strip()
    if key in cover_map:
        c = row.cells[1]
        # 셀 첫 단락 텍스트 교체(서식 유지)
        p = c.paragraphs[0]
        if p.runs:
            p.runs[0].text = cover_map[key]
            for r in p.runs[1:]:
                r._r.getparent().remove(r._r)
        else:
            p.add_run(cover_map[key])

# ── 2. 본문 제거 (표지 정보표 이후 ~ sectPr 직전) ──────
body = doc.element.body
children = list(body.iterchildren())
# 표지 정보표(tbl) 인덱스 찾기
cover_tbl_el = cover._tbl
start_idx = None
for i, c in enumerate(children):
    if c is cover_tbl_el:
        start_idx = i
        break
assert start_idx is not None
# 정보표 다음부터 sectPr 전까지 제거
for c in children[start_idx + 1:]:
    tag = c.tag.split('}')[-1]
    if tag == 'sectPr':
        break
    body.remove(c)

# 이후 append 기준점: sectPr 앞에 삽입해야 함
sectPr = body.find(qn('w:sectPr'))

# ── 헬퍼: sectPr 앞에 단락/표 추가 ───────────────────
def _new_p():
    p = OxmlElement('w:p')
    sectPr.addprevious(p)
    return Paragraph(p, doc)

def _set_run_font(run, size=10.5, bold=False, color=C_TEXT, mono=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    name = FONT_MONO if mono else FONT_KO
    run.font.name = name
    rpr = run._r.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    rf.set(qn('w:ascii'), name); rf.set(qn('w:hAnsi'), name); rf.set(qn('w:eastAsia'), name)

INLINE_RE = re.compile(r'(\*\*.+?\*\*|`[^`]+`)')

def add_runs_with_inline(p, text, base_size=10.5, base_color=C_TEXT):
    """**bold**, `code` 인라인 처리."""
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2]); _set_run_font(r, base_size, True, base_color)
        elif part.startswith('`') and part.endswith('`'):
            r = p.add_run(part[1:-1]); _set_run_font(r, base_size - 0.5, False, RGBColor(0xB0,0x2A,0x37), mono=True)
        else:
            r = p.add_run(part); _set_run_font(r, base_size, False, base_color)

def add_heading2(num, title):
    p = _new_p()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r1 = p.add_run(num + " "); _set_run_font(r1, 15, True, C_NUM)
    r2 = p.add_run(title); _set_run_font(r2, 15, True, C_TITLE)

def add_heading3(num, title):
    p = _new_p()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r0 = p.add_run("▍"); _set_run_font(r0, 12, False, C_NUM)
    r1 = p.add_run(" " + num + " " + title); _set_run_font(r1, 12, True, C_TITLE)

def add_heading4(title):
    p = _new_p()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title); _set_run_font(r, 11, True, C_TITLE)

def add_body(text):
    p = _new_p()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    add_runs_with_inline(p, text)

def add_bullet(text, level=0):
    p = _new_p()
    p.paragraph_format.left_indent = Pt(18 + level*14)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.25
    b = p.add_run("• "); _set_run_font(b, 10.5, False, C_TITLE)
    add_runs_with_inline(p, text)

def add_code_block(lines, caption=None):
    if caption:
        cp = _new_p(); cp.paragraph_format.space_before = Pt(6); cp.paragraph_format.space_after = Pt(2)
        r = cp.add_run(caption); _set_run_font(r, 9.5, True, C_TITLE)
    p = _new_p()
    p.paragraph_format.space_after = Pt(6)
    # 음영 박스
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:fill'),'F4F6FB')
    pPr.append(shd)
    bdr = OxmlElement('w:pBdr')
    for side in ('top','bottom','left','right'):
        e = OxmlElement('w:'+side); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'4'); e.set(qn('w:space'),'4'); e.set(qn('w:color'),'D5DCEC')
        bdr.append(e)
    pPr.append(bdr)
    for i, ln in enumerate(lines):
        if i > 0:
            br = p.add_run(); br.add_break()
        r = p.add_run(ln if ln else " "); _set_run_font(r, 9, False, RGBColor(0x1F,0x29,0x37), mono=True)

def _shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:fill'),fill)
    tcPr.append(shd)

def _set_cell_text(cell, text, bold=False, color=C_TEXT, size=9.5, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)
    add_runs_with_inline(p, text, base_size=size, base_color=color)
    if bold:
        for r in p.runs: r.font.bold = True

def _set_table_borders(tbl):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e = OxmlElement('w:'+edge)
        e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'4'); e.set(qn('w:space'),'0'); e.set(qn('w:color'),'C9D3E6')
        borders.append(e)
    tblPr.append(borders)

def add_table(header, rows):
    ncol = len(header)
    tbl = doc.add_table(rows=1, cols=ncol)
    _set_table_borders(tbl)
    # add_table는 문서 끝에 붙으므로 sectPr 앞으로 이동
    sectPr.addprevious(tbl._tbl)
    # 헤더
    for j, h in enumerate(header):
        c = tbl.rows[0].cells[j]
        _shade_cell(c, HDR_FILL)
        _set_cell_text(c, h, bold=True, color=RGBColor(0xFF,0xFF,0xFF), size=9.5)
    # 행
    for ri, row in enumerate(rows):
        cells = tbl.add_row().cells
        for j in range(ncol):
            val = row[j] if j < len(row) else ""
            c = cells[j]
            if ri % 2 == 1:
                _shade_cell(c, ROW_ALT)
            _set_cell_text(c, val, size=9.5)
    # 간격
    sp = _new_p(); sp.paragraph_format.space_after = Pt(2)

def add_image(path, caption=None, width_in=5.5):
    if not os.path.exists(path):
        add_body(f"[이미지 누락: {os.path.basename(path)}]")
        return
    p = _new_p(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Emu(int(width_in * 914400)))
    if caption:
        cp = _new_p(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER; cp.paragraph_format.space_after = Pt(8)
        r = cp.add_run(caption); _set_run_font(r, 9, False, RGBColor(0x64,0x74,0x8B))

# ── 3. 마크다운 파싱 → 렌더 ──────────────────────────
with open(MD, encoding='utf-8') as f:
    raw = f.read()

# 릴리즈 버전은 본문(final-report.md §6.1)에서 client-v0.1.2 기준으로 이미 정리됨.
lines = raw.split('\n')
i = 0
n = len(lines)
skip_until_first_h2 = True   # 표지/목차는 템플릿 표지로 대체 → 첫 '## 1.' 전까지 스킵

def parse_table(start):
    """마크다운 표 파싱. returns (header, rows, next_index)."""
    header = [c.strip() for c in lines[start].strip().strip('|').split('|')]
    j = start + 2  # 구분선 건너뜀
    rows = []
    while j < n and lines[j].strip().startswith('|'):
        cells = [c.strip() for c in lines[j].strip().strip('|').split('|')]
        rows.append(cells)
        j += 1
    return header, rows, j

while i < n:
    line = lines[i]
    s = line.strip()

    # 스킵 모드: 첫 '## 1.' 만나기 전까지 표지/목차 무시
    if skip_until_first_h2:
        m = re.match(r'^##\s+(\d+)\.\s+(.+)$', s)
        if m:
            skip_until_first_h2 = False
        else:
            i += 1
            continue

    # 코드블록
    if s.startswith('```'):
        lang = s[3:].strip()
        block = []
        i += 1
        while i < n and not lines[i].strip().startswith('```'):
            block.append(lines[i]); i += 1
        i += 1  # 닫는 ```
        if lang == 'mermaid':
            add_code_block(block, caption="[그림] 다이어그램 (Mermaid 소스)")
        else:
            add_code_block(block)
        continue

    # 헤딩
    m2 = re.match(r'^##\s+(\d+)\.\s+(.+)$', s)
    m3 = re.match(r'^###\s+([\d.]+)\s+(.+)$', s)
    m3b = re.match(r'^###\s+(.+)$', s)
    if m2:
        add_heading2(m2.group(1) + ".", m2.group(2)); i += 1; continue
    if m3:
        add_heading3(m3.group(1), m3.group(2)); i += 1; continue
    if s.startswith('#### '):
        add_heading4(s[5:].strip()); i += 1; continue
    if m3b and s.startswith('### '):
        add_heading4(m3b.group(1)); i += 1; continue

    # 수평선 / 빈 줄
    if s in ('---', ''):
        i += 1; continue

    # '## 참조 문서' 등 번호 없는 H2
    mh2plain = re.match(r'^##\s+(.+)$', s)
    if mh2plain:
        add_heading2("", mh2plain.group(1)); i += 1; continue

    # 표
    if s.startswith('|') and i + 1 < n and re.match(r'^\|?[\s:|-]+\|', lines[i+1].strip()):
        header, rows, j = parse_table(i)
        # 표 안에 이미지가 있으면 이미지 그리드로 렌더
        flat = " ".join(" ".join(r) for r in rows)
        if re.search(r'!\[.*?\]\(.+?\)', flat):
            img_re = re.compile(r'!\[(.*?)\]\((.+?)\)')
            for r in rows:
                for cell in r:
                    mm = img_re.search(cell)
                    if mm:
                        alt, path = mm.group(1), mm.group(2)
                        abspath = os.path.normpath(os.path.join(os.path.dirname(MD), path))
                        # 한 행에 여러 이미지면 폭 축소
                        ncell_imgs = sum(1 for cc in r if img_re.search(cc))
                        w = 5.3 if ncell_imgs == 1 else (4.5 if ncell_imgs == 2 else 3.4)
                        add_image(abspath, caption=alt or None, width_in=w)
        else:
            add_table(header, rows)
        i = j; continue

    # 이미지 ![alt](path)
    mimg = re.match(r'^!\[(.*?)\]\((.+?)\)$', s)
    if mimg:
        alt, path = mimg.group(1), mimg.group(2)
        abspath = os.path.normpath(os.path.join(os.path.dirname(MD), path))
        add_image(abspath, caption=alt or None)
        i += 1; continue

    # 리스트
    mbul = re.match(r'^(\s*)[-*]\s+(.+)$', line)
    if mbul:
        indent = len(mbul.group(1))
        add_bullet(mbul.group(2), level=1 if indent >= 2 else 0)
        i += 1; continue
    mnum = re.match(r'^\s*\d+\.\s+(.+)$', line)
    if mnum:
        add_bullet(mnum.group(1))
        i += 1; continue

    # 일반 본문
    add_body(s)
    i += 1

# ── 4. 저장 ──────────────────────────────────────────
os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print("saved:", OUT)
print("paragraphs:", len(doc.paragraphs), "| tables:", len(doc.tables))
